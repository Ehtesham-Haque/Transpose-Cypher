import tkinter as tk
from tkinter import messagebox, filedialog
from docx import Document
import os
import numpy as np
import re

# Encryption and decryption functions (same as your original code)
def create_matrix(rows, cols, char_sequence):
    padded_sequence = char_sequence.ljust(rows * cols, '#')
    A = np.zeros((rows, cols), dtype=int)
    seq_index = 0
    for i in range(rows):
        for j in range(cols):
            if seq_index < len(padded_sequence):
                A[i, j] = ord(padded_sequence[seq_index].upper())
                seq_index += 1
    return A

def transpose_matrix(A):
    B = A.T
    return B

def compress_hashes(s):
    return re.sub(r'(#+)', lambda m: str(len(m.group(0))) + '#', s)

def decompress_hashes(s):
    return re.sub(r'(\d+)#', lambda m: '#' * int(m.group(1)), s)

def encrypt_message(message, rows, cols):
    A = create_matrix(rows, cols, message)
    B = transpose_matrix(A)
    transposed_line = ''.join(chr(value) for value in B.flatten() if value != 0)
    compressed_message = compress_hashes(transposed_line)
    return compressed_message

def decrypt_message(encrypted_message, rows, cols):
    expanded_message = decompress_hashes(encrypted_message)
    encrypted_length = len(expanded_message)
    B = np.zeros((cols, rows), dtype=int)
    for i in range(encrypted_length):
        B[i // rows, i % rows] = ord(expanded_message[i])
    A = B.T.flatten()[:encrypted_length]
    original_message = ''.join(chr(value) for value in A)
    return original_message.replace('#', ' ')

# Tkinter GUI setup
class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Transpose Cypher")
        self.root.configure(bg="#1f1f1f")  # Set background color of the root window to dark gray

        # Top Frame for row and column entry fields
        self.top_frame = tk.Frame(self.root, bg="#1f1f1f")  # Set background color of top frame to dark gray
        self.top_frame.pack(pady=20)

        self.rows_label = tk.Label(self.top_frame, text="Enter Row [m]", fg="white", bg="#1f1f1f", font=('Helvetica', 15, 'bold'))
        self.rows_label.grid(row=0, column=0, padx=10, pady=5)

        self.rows_entry = tk.Entry(self.top_frame,bg="#8d8d8d",font=('Helvetica', 20))
        self.rows_entry.grid(row=0, column=1, padx=10, pady=5)

        self.cols_label = tk.Label(self.top_frame, text="Enter Column [n]", fg="white", bg="#1f1f1f", font=('Helvetica', 15, 'bold'))
        self.cols_label.grid(row=1, column=0, padx=10, pady=5)

        self.cols_entry = tk.Entry(self.top_frame,bg="#8d8d8d", font=('Helvetica', 20))
        self.cols_entry.grid(row=1, column=1, padx=10, pady=5)

        # Bottom Frame for encrypt, decrypt, and exit buttons
        self.bottom_frame = tk.Frame(self.root, bg="#1f1f1f")  # Set background color of bottom frame to dark gray
        self.bottom_frame.pack()

        # Function to create styled buttons with larger text
        def create_button(parent, text, command):
            button = tk.Button(parent, text=text, fg="black", bg="white", font=('Helvetica', 15, 'bold'), command=command)
            return button

        self.encrypt_button = create_button(self.bottom_frame, "Encrypt", self.encrypt)
        self.encrypt_button.grid(row=0, column=0, padx=10, pady=10)

        self.decrypt_button = create_button(self.bottom_frame, "Decrypt", self.decrypt)
        self.decrypt_button.grid(row=0, column=1, padx=10, pady=10)

        self.exit_button = create_button(self.bottom_frame, "Exit", root.quit)
        self.exit_button.grid(row=0, column=2, padx=10, pady=10)

    def encrypt(self):
        # Select Plaintext.docx file
        plaintext_path = filedialog.askopenfilename(title="Select Plaintext.docx",
                                                    filetypes=[("Word files", "*.docx")])
        if not plaintext_path:
            return

        # Read plaintext from selected file
        doc = Document(plaintext_path)
        message = '\n'.join([para.text for para in doc.paragraphs])

        # Get matrix size from user input
        try:
            rows = int(self.rows_entry.get())
            cols = int(self.cols_entry.get())
        except ValueError:
            messagebox.showerror("Error", "Please enter valid integers for rows and columns.")
            return

        # Encrypt the message
        encrypted_message = encrypt_message(message, rows, cols)

        # Save encrypted message to Encrypted text.docx
        encrypted_path = os.path.join(os.path.dirname(plaintext_path), "Encrypted text.docx")
        doc = Document()
        doc.add_paragraph(encrypted_message)
        doc.save(encrypted_path)
        messagebox.showinfo("Encryption", "Message encrypted and saved to Encrypted text.docx")

    def decrypt(self):
        # Select Encrypted text.docx file
        encrypted_path = filedialog.askopenfilename(title="Select Encrypted text.docx",
                                                    filetypes=[("Word files", "*.docx")])
        if not encrypted_path:
            return

        # Read encrypted message from selected file
        doc = Document(encrypted_path)
        encrypted_message = '\n'.join([para.text for para in doc.paragraphs])

        # Get matrix size from user input
        try:
            rows = int(self.rows_entry.get())
            cols = int(self.cols_entry.get())
        except ValueError:
            messagebox.showerror("Error", "Please enter valid integers for rows and columns.")
            return

        # Decrypt the message
        decrypted_message = decrypt_message(encrypted_message, rows, cols)

        # Save decrypted message to Decrypted text.docx
        decrypted_path = os.path.join(os.path.dirname(encrypted_path), "Decrypted text.docx")
        doc = Document()
        doc.add_paragraph(decrypted_message)
        doc.save(decrypted_path)
        messagebox.showinfo("Decryption", "Message decrypted and saved to Decrypted text.docx")

# Main function to start the GUI
def main():
    root = tk.Tk()
    app = App(root)
    root.mainloop()

if __name__ == "__main__":
    main()
